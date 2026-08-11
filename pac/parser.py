#!/usr/bin/env python3
"""AI采购助理 - 多格式报价解析：CSV / 邮件文本TXT / Excel XLSX -> SupplierQuote"""
from __future__ import annotations
import csv
import os
import re
from .models import LineItem, SupplierQuote

# 常见中文表头 -> 标准字段
HEADER_MAP = {
    "品名": "item", "物料名称": "item", "产品名称": "item", "名称": "item",
    "规格": "spec", "型号": "spec",
    "数量": "qty", "采购数量": "qty", "需求数量": "qty",
    "单价": "unit_price", "含税单价": "unit_price", "报价": "unit_price", "价格": "unit_price",
    "起订量": "moq", "最小起订量": "moq",
    "交期": "lead_time", "交货期": "lead_time", "交期(天)": "lead_time",
    "账期": "payment", "付款条件": "payment", "付款方式": "payment",
    "运费": "freight", "物流费": "freight",
    "有效期": "validity", "报价有效期": "validity",
}


def _norm_header(h: str) -> str:
    h = h.strip().lower()
    if h in HEADER_MAP:
        return HEADER_MAP[h]
    # 模糊匹配真实世界表头（如"原料名称""报价（元）""交货周期（天）"）
    if "供应商" in h:
        pass  # 供应商名列不映射到item
    elif "名称" in h or "品名" in h or "物料" in h:
        return "item"
    if "规格" in h or "型号" in h:
        return "spec"
    if "数量" in h or "采购量" in h:
        return "qty"
    if "报价" in h or "价格" in h or "单价" in h:
        return "unit_price"
    if "交期" in h or "交货" in h or "周期" in h:
        return "lead_time"
    if "付款" in h or "账期" in h or "月结" in h:
        return "payment"
    if "起订" in h:
        return "moq"
    if "运费" in h or "物流" in h:
        return "freight"
    if "有效" in h:
        return "validity"
    return h


def parse_csv(path: str) -> list[SupplierQuote]:
    """解析结构化CSV：第一行为表头（中文亦可）。返回列表（单元素）。"""
    with open(path, encoding="utf-8-sig", errors="replace") as f:
        rows = list(csv.DictReader(f))
    headers = list(rows[0].keys()) if rows else []
    cols = {_norm_header(h): h for h in headers}
    supplier = os.path.splitext(os.path.basename(path))[0]
    q = SupplierQuote(supplier=supplier, source_file=os.path.basename(path))
    for r in rows:
        try:
            item = r.get(cols.get("item", ""), "")
            spec = r.get(cols.get("spec", ""), "")
            qty = int(float(str(r.get(cols.get("qty", ""), "0")).replace(",", "")))
            price = float(str(r.get(cols.get("unit_price", ""), "0")).replace(",", ""))
            if item and price > 0:
                q.items.append(LineItem(item=item, spec=spec, qty=qty, unit_price=price))
        except (ValueError, TypeError):
            continue
    # 附加字段（若表中有）
    for k, v in [("moq", "moq"), ("lead_time", "交期"), ("payment", "账期"),
                 ("freight", "运费"), ("validity", "有效期")]:
        if k in cols and rows and rows[0].get(cols[k]):
            val = rows[0][cols[k]]
            if k == "freight":
                if "含" in val:
                    q.freight_inc = True
                else:
                    try:
                        q.freight = float(val)
                    except ValueError:
                        pass
            else:
                setattr(q, k, val.strip())
    return [q]


def parse_xlsx(path: str) -> list[SupplierQuote]:
    """解析Excel报价单（pandas + openpyxl）。
    若表中有"供应商名称"列（市场调研表，一料多商），按供应商拆分为多个SupplierQuote。
    返回列表。"""
    import pandas as pd
    supplier = os.path.splitext(os.path.basename(path))[0]
    df = pd.read_excel(path, sheet_name=None, dtype=str)
    quotes: list[SupplierQuote] = []
    for sheet, data in df.items():
        data = data.dropna(how="all")
        if data.empty:
            continue
        headers = [str(c).strip() for c in data.columns]
        cols = {_norm_header(h): h for h in headers}
        if "item" not in cols and "unit_price" not in cols:
            continue
        # 检测供应商列（一料多商结构）
        sup_col = next((h for h in headers if "供应商" in h), None)

        def make_quote(name: str) -> SupplierQuote:
            q = SupplierQuote(supplier=name, source_file=os.path.basename(path))
            for k, v in [("moq", "moq"), ("lead_time", "交期"), ("payment", "账期"),
                         ("freight", "运费"), ("validity", "有效期")]:
                if k in cols and len(data) and data.iloc[0].get(cols[k]) is not None:
                    val = str(data.iloc[0][cols[k]]).strip()
                    if k == "freight":
                        if "含" in val:
                            q.freight_inc = True
                        else:
                            try:
                                q.freight = float(val)
                            except ValueError:
                                pass
                    else:
                        setattr(q, k, val)
            return q

        groups: dict[str, list[LineItem]] = {}
        for _, r in data.iterrows():
            try:
                item = str(r.get(cols.get("item", ""), "") or "")
                spec = str(r.get(cols.get("spec", ""), "") or "")
                qty = int(float(str(r.get(cols.get("qty", ""), "0")).replace(",", "")))
                price = float(str(r.get(cols.get("unit_price", ""), "0")).replace(",", ""))
                if item and item != "nan" and price > 0:
                    # 数量缺失（市场报价调研表）时按单价比较模式处理
                    if qty <= 0:
                        qty = 1
                    key = str(r.get(sup_col, "") or "").strip() if sup_col else supplier
                    if not key or key == "nan":
                        key = supplier
                    groups.setdefault(key, []).append(
                        LineItem(item=item, spec=spec, qty=qty, unit_price=price))
            except (ValueError, TypeError):
                continue
        if not groups:
            continue
        for name, items in groups.items():
            q = make_quote(name)
            q.items = items
            quotes.append(q)
        if quotes:
            break
    return quotes


def parse_docx(path: str) -> list[SupplierQuote]:
    """解析Word采购/报价记录：标题段定物料名，'采购N桶，单价X元'行提取
    同一物料多条记录（月度）保留文档中最后一条（假定按时间升序=最新）"""
    from docx import Document
    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                texts.append(" | ".join(cells))
    text = "\n".join(texts)

    q = SupplierQuote(supplier=os.path.splitext(os.path.basename(path))[0],
                      source_file=os.path.basename(path), raw_text=text)
    _extract_terms(q, text)

    # 标题段: "1. 可可脂，规格25kg/桶，单位：桶"
    title_re = re.compile(r"^\d+\.\s*([\u4e00-\u9fa5A-Za-z]{2,15})[，,]\s*规格[:：]?\s*([^\s，,]+)")
    # 采购行: "2025年10月：采购120桶，单价185元/桶，金额22200元，供货来源XX公司；"
    buy_re = re.compile(r"采购\s*(\d+)\s*(桶|袋|个|只|件|套|吨|kg)\s*[，,]\s*单价\s*([1-9]\d*(?:\.\d+)?)\s*元")
    # 供货来源: "供货来源嘉吉投资（中国）有限公司"
    sup_re = re.compile(r"供货来源\s*([\u4e00-\u9fa5A-Za-z0-9（）()]{2,30}公司)")

    current_item, current_spec = "", ""
    last_by_key: dict[tuple, LineItem] = {}
    for line in texts:
        mt = title_re.match(line)
        if mt:
            current_item, current_spec = mt.group(1), mt.group(2)
            continue
        mb = buy_re.search(line)
        if mb and current_item:
            ms = sup_re.search(line)
            src = ms.group(1) if ms else ""
            key = (current_item, current_spec)
            last_by_key[key] = LineItem(
                item=current_item, spec=current_spec,
                qty=int(mb.group(1)), unit_price=float(mb.group(3)),
                note=f"历史采购价（{src}）",
            )
    q.items = list(last_by_key.values())
    return [q]


def _extract_terms(q: SupplierQuote, text: str):
    """从文本抽取交易条款（供txt/docx共用）"""
    m = re.search(r"起订量[量为:：]?\s*(\d+[万k个件]*)", text)
    if m:
        q.moq = m.group(1)
    m = re.search(r"交期[为:：]?\s*(\d+[-~至]?\d*\s*天)", text)
    if m:
        q.lead_time = m.group(1)
    m = re.search(r"(?:账期|月结|付款方式)[为:：]?\s*([^\n，。;；]{2,20})", text)
    if m:
        q.payment = m.group(1).strip()
    m = re.search(r"运费[为:：]?\s*(?:约|大概)?\s*(\d+(?:\.\d+)?)\s*元", text)
    if m:
        q.freight = float(m.group(1))
    m = re.search(r"有效期[为:：]?\s*([^\n。;；]{2,15})", text)
    if m:
        q.validity = m.group(1).strip()


def _extract_items(q: SupplierQuote, text: str):
    """从文本抽取物料行（供txt/docx共用）"""
    blacklist = ("起订", "交期", "运费", "有效期", "报价", "含税", "单价", "数量")
    for m in re.finditer(
        r"([\u4e00-\u9fa5A-Za-z]{2,20}?)\s*"
        r"([0-9*×xX\-–—.]{3,20}cm|[\u4e00-\u9fa5]{1,6}克|[\u4e00-\u9fa5]{1,6}层|[\u4e00-\u9fa5]{1,6}kg)?\s*"
        r"(\d[\d,]*)\s*(?:个|只|件|套|桶|袋|pcs|PCS)?\s*"
        r"(?:单价|价格|报价)?[为:：]?\s*([1-9]\d*(?:\.\d+)?)\s*元",
        text,
    ):
        item, spec, qty, price = m.group(1), m.group(2) or "", m.group(3), m.group(4)
        if any(b in item for b in blacklist):
            continue
        q.items.append(LineItem(
            item=item, spec=spec,
            qty=int(qty.replace(",", "")),
            unit_price=float(price),
        ))


def parse_txt_email(path: str) -> SupplierQuote:
    """解析非结构化邮件文本：正则抓供应商名、物料行、交易条款"""
    with open(path, encoding="utf-8-sig", errors="replace") as f:
        text = f.read()
    supplier = os.path.splitext(os.path.basename(path))[0]
    # 去掉邮件头
    q = SupplierQuote(supplier=supplier, source_file=os.path.basename(path), raw_text=text)

    # 供应商名（"XX公司"或"XX实业"等）——取最长匹配，避免匹配到主题行短名
    names = re.findall(r"([\u4e00-\u9fa5A-Za-z0-9]{2,15}(?:公司|实业|包装|纸业|集团|厂))", text)
    if names:
        q.supplier = max(names, key=len)

    _extract_terms(q, text)
    _extract_items(q, text)
    return [q]


def parse_any(path: str) -> list[SupplierQuote]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return parse_csv(path)
    if ext == ".xlsx":
        return parse_xlsx(path)
    if ext in (".txt", ".md", ".eml"):
        return parse_txt_email(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext == ".pdf":
        from .pdf_parser import parse_pdf
        return [parse_pdf(path)]
    raise ValueError(f"不支持的文件格式: {ext}")
